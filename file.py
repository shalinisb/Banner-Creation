from dotenv import load_dotenv
from docx import Document
from PIL import Image
import fitz,pandas as pd,io
import pytesseract
load_dotenv()

def parse_table_from_docx(doc):
    tables_data = []
    for table_num, table in enumerate(doc.tables):
        table_data = {
            f"Table {table_num + 1}": []
        }

        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data[f"Table {table_num + 1}"].append(row_data)

        tables_data.append(table_data)

    return str(tables_data)


def extract_text_from_pdf(pdf_file):
    text = ""
    with fitz.open("pdf", pdf_file.file.read()) as pdf_doc:
        for page_num in range(pdf_doc.page_count):
            page = pdf_doc[page_num]
            text += page.get_text()
    return text


def extract_text_from_docx(docx_file):
    doc = Document(io.BytesIO(docx_file.file.read()))
    text = ""
    # Paragraph Data
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    # Extract Table Data
    tables_data = parse_table_from_docx(doc)
    return text+tables_data

def extract_text_from_txt(txt_file):
    with io.TextIOWrapper(io.BytesIO(txt_file.file.read()), encoding='utf-8') as file:
        text = file.read()
    return text

def extract_text_from_excel(excel_file):
    df = pd.read_excel(io.BytesIO(excel_file.file.read()))
    text = df.to_string(index=False)
    return text



def extract_text_from_image(image_file):
    pytesseract.pytesseract.tesseract_cmd = r"./tesseract/tesseract.exe"
    image = Image.open(io.BytesIO(image_file.file.read()))
    text = pytesseract.image_to_string(image)
    return text


# Process File
class FileProcessor():
    def __init__(self,file):
            self.file = file
            self.filename = file.filename
            self.content_type = file.content_type

    def process_file_to_txt(self):
            if self.filename.split('.')[-1] == 'pdf':
                return extract_text_from_pdf(self.file)
            elif self.filename.split('.')[-1] == 'docx':
                return extract_text_from_docx(self.file)
            elif self.filename.split('.')[-1] == 'txt':
                return extract_text_from_txt(self.file)
            elif self.filename.split('.')[-1] == 'xlsx':
                return extract_text_from_excel(self.file)
            elif self.filename.split('.')[-1] in ('jpg', 'jpeg', 'png', 'gif', 'bmp'):
                return extract_text_from_image(self.file)
            else:
                return "Unsupported file type"

class FileBlobWrapper():
    def __init__(self,blob,filename,content_type):
        self.file = blob
        self.filename = filename
        self.content_type = content_type
